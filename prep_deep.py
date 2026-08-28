import docx
import os

def prep_deep(input_file, output_file):
    doc = docx.Document(input_file)
    
    # 1. Replace specific cover strings
    for p in doc.paragraphs:
        if "Judul Penelitian(Times New Roman 14, Bold)" in p.text:
            p.text = p.text.replace("Judul Penelitian(Times New Roman 14, Bold)", "${JUDUL}")
        elif "Judul Penelitian" in p.text and len(p.text) < 25 and not "(Tulis" in p.text:
            p.text = p.text.replace("Judul Penelitian", "${JUDUL}")
            
        if "Nama Pengusul, lengkap dengan gelar" in p.text:
            p.text = p.text.replace("Nama Pengusul, lengkap dengan gelar (NIDN)", "${NAMA_KETUA}")
        if "Nama Peneliti, lengkap dengan gelar" in p.text:
            p.text = p.text.replace("Nama Peneliti, lengkap dengan gelar (NIDN)", "${NAMA_KETUA}")
            
        if "Nama Anggota(Times New Roman 12, Bold)" in p.text:
            p.text = p.text.replace("Nama Anggota(Times New Roman 12, Bold)", "${NAMA_ANGGOTA}")
        elif "Nama Anggota" in p.text and len(p.text) < 20:
            p.text = p.text.replace("Nama Anggota", "${NAMA_ANGGOTA}")

    # 2. Section Replacements
    # We will look for headers like "RINGKASAN", "PENDAHULUAN", "METODE PENELITIAN"
    # and replace the content of the paragraph immediately following it.
    
    def replace_after_heading(heading_text, placeholder):
        found = False
        for i, p in enumerate(doc.paragraphs):
            if heading_text in p.text.upper():
                # Look at the next few paragraphs to insert placeholder
                for j in range(i+1, min(i+5, len(doc.paragraphs))):
                    text = doc.paragraphs[j].text.strip()
                    if text and not text.startswith('BAB') and not text.startswith('Tabel') and not text.startswith('Gambar'):
                        # Found the first content paragraph
                        doc.paragraphs[j].text = placeholder
                        return

    replace_after_heading("RINGKASAN(", "${RINGKASAN}") # "RINGKASAN(Maksimum satu halaman)"
    replace_after_heading("BAB I PENDAHULUAN", "${PENDAHULUAN}")
    replace_after_heading("METODE PENELITIAN", "${METODE}")
    replace_after_heading("DAFTAR PUSTAKA", "${PUSTAKA}")
    replace_after_heading("HASIL DAN LUARAN YANG DICAPAI", "${LUARAN}")
    replace_after_heading("LUARAN DAN TARGET CAPAIAN", "${LUARAN}")

    # 3. Budget Tables
    # We want to replace the rows inside the budget tables with our own cloned rows.
    # We will identify the tables by their headers.
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
            first_cell = table.rows[0].cells[0].text.strip()
            if first_cell == "No":
                # Check if it's the budget table
                if len(table.rows[0].cells) >= 4 and "Total (Rp)" in table.rows[0].cells[-1].text:
                    # We will clear all rows except header and total, and add a placeholder row for cloneRow
                    # Actually, PhpWord cloneRow needs a placeholder in a row.
                    # Let's just put ${B_NO}, ${B_ITEM}, ${B_VOL}, ${B_HARGA}, ${B_TOTAL}
                    if len(table.rows) > 2:
                        table.rows[1].cells[0].text = "${B_NO}"
                        table.rows[1].cells[1].text = "${B_ITEM}"
                        table.rows[1].cells[2].text = "1"
                        table.rows[1].cells[3].text = "${B_HARGA}"
                        table.rows[1].cells[4].text = "${B_TOTAL}"
                        
                        # Set total
                        table.rows[-1].cells[-1].text = "${B_GRAND}"
                        
    # 4. Jadwal Penelitian
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
            first_cell = table.rows[0].cells[0].text.strip()
            if first_cell == "No" and "Nama Kegiatan" in table.rows[0].cells[1].text:
                if len(table.rows) > 1:
                    table.rows[1].cells[0].text = "${J_NO}"
                    table.rows[1].cells[1].text = "${J_KEGIATAN}"
                    for c in range(2, len(table.rows[1].cells)):
                        table.rows[1].cells[c].text = "v"

    doc.save(output_file)

# Since we already patched the placeholders like (Tulis ...), we can run this on the already prepped templates!
prep_deep("storage/app/templates/Template_Laporan.docx", "storage/app/templates/Template_Laporan.docx")
prep_deep("storage/app/templates/Template_Proposal.docx", "storage/app/templates/Template_Proposal.docx")
print("Deep prep completed.")
