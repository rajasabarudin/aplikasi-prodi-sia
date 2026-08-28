import docx
import os
import re

def prep_template(input_file, output_file):
    doc = docx.Document(input_file)
    
    replacements = {
        "(Tulis judul penelitian)": "${JUDUL}",
        "BULAN TAHUN": "${BULAN_TAHUN}",
        
        "(Tulis nama ketua Peneliti beserta gelar)": "${NAMA_KETUA}",
        "(Tulis nama ketua pengusul beserta gelar)": "${NAMA_KETUA}",
        
        "(Tulis NIDN ketua Peneliti)": "${NIDN_KETUA}",
        "(Tulis NIDN ketua pengusul)": "${NIDN_KETUA}",
        
        "(Tulis Jabatan Fungsional ketua Peneliti)": "${JABATAN_KETUA}",
        "(Tulis Jabatan Fungsional ketua pengusul)": "${JABATAN_KETUA}",
        
        "(Tulis Program Studi ketua Peneliti)": "${PRODI_KETUA}",
        "(Tulis Program Studi ketua pengusul)": "${PRODI_KETUA}",
        
        "(Tulis HP ketua Peneliti)": "${HP_KETUA}",
        "(Tulis HP ketua pengusul)": "${HP_KETUA}",
        
        "(Tulis surel ketua Peneliti)": "${EMAIL_KETUA}",
        "(Tulis surel ketua pengusul)": "${EMAIL_KETUA}",
        
        "(Tulis nama anggota Peneliti)": "${NAMA_ANGGOTA}",
        "(Tulis nama anggota pengusul)": "${NAMA_ANGGOTA}",
        
        "(Tulis NIDN anggota Peneliti)": "${NIDN_ANGGOTA}",
        "(Tulis NIDN anggota pengusul)": "${NIDN_ANGGOTA}",
        "(Tulis anggota anggota Peneliti)": "${NIDN_ANGGOTA}",
        "(Tulis anggota anggota pengusul)": "${NIDN_ANGGOTA}",
        
        "(Tulis Jabatan Fungsional anggota Peneliti)": "${JABATAN_ANGGOTA}",
        "(Tulis Jabatan Fungsional anggota pengusul)": "${JABATAN_ANGGOTA}",
        
        "(Tulis Program Studi anggota Peneliti)": "${PRODI_ANGGOTA}",
        "(Tulis Program Studi anggota pengusul)": "${PRODI_ANGGOTA}",
        
        "(Tulis nama mitra jika ada)": "${NAMA_MITRA}",
        "(Tulis alamat mitra jika ada)": "${ALAMAT_MITRA}",
        "(Tulis penanggung jawab)": "${PJ_MITRA}",
        
        "(Nama Ketua LPPM beserta gelar)": "${KETUA_LPPM}",
        "(Nama Ketua beserta gelar)": "${KETUA_LPPM}",
        "(Nama Rektor beserta gelar)": "${REKTOR}",
        
        "(Tulis nama dan gelar Pengusul)": "${NAMA_KETUA}",
        "(Tulis NIDN Pengusul)": "${NIDN_KETUA}",
        "(Tulis Jabatan Fungsional Pengusul)": "${JABATAN_KETUA}",
        "(Tulis Program Studi Pengusul)": "${PRODI_KETUA}",
        "(Tulis institusi UBSI/PSDKU)": "${INSTITUSI}",
        "(Tulis jangka waktu penelitian)": "${WAKTU_PENELITIAN}",
        "(Nama Pengusul)": "${NAMA_KETUA}"
    }

    def replace_text(p):
        for key, val in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, val)
        # Handle special cases for Biaya
        if "Biaya yang disetujui: Rp" in p.text:
            p.text = "Biaya yang disetujui: Rp ${BIAYA}"
        if "Biaya yang diusulkan: Rp" in p.text:
            p.text = "Biaya yang diusulkan: Rp ${BIAYA}"
        # Titles
        if "Judul Penelitian(Times New Roman 14, Bold)" in p.text:
            p.text = p.text.replace("Judul Penelitian(Times New Roman 14, Bold)", "${JUDUL}")
        
    for p in doc.paragraphs:
        replace_text(p)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text(p)
                        
    doc.save(output_file)

os.makedirs("storage/app/templates", exist_ok=True)

# Process from the original templates to avoid double-processing already mangled text
prep_template("Template Laporan Penelitian Mandiri.docx", "storage/app/templates/Template_Laporan.docx")
prep_template("Template Proposal penelitian mandiri - Copy.docx", "storage/app/templates/Template_Proposal.docx")
print("Templates prepped with all variables.")
