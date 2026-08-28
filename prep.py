import docx
import os

def prep_template(input_file, output_file):
    doc = docx.Document(input_file)
    
    # Replace in paragraphs
    for p in doc.paragraphs:
        if "(Tulis judul penelitian)" in p.text:
            p.text = p.text.replace("(Tulis judul penelitian)", "${JUDUL}")
        if "BULAN TAHUN" in p.text:
            p.text = p.text.replace("BULAN TAHUN", "${BULAN_TAHUN}")
        # There is a line: "Judul Penelitian(Times New Roman 14, Bold)"
        # Let's just find "Judul Penelitian" that is alone or similar and replace it,
        # But wait, we can just replace specific strings.
        if "Judul Penelitian(Times New Roman 14, Bold)" in p.text:
            p.text = p.text.replace("Judul Penelitian(Times New Roman 14, Bold)", "${JUDUL}")
            
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "(Tulis judul penelitian)" in p.text:
                        p.text = p.text.replace("(Tulis judul penelitian)", "${JUDUL}")
                    if "BULAN TAHUN" in p.text:
                        p.text = p.text.replace("BULAN TAHUN", "${BULAN_TAHUN}")
                        
    doc.save(output_file)

# Ensure the template directory exists
os.makedirs("storage/app/templates", exist_ok=True)

prep_template("Template Laporan Penelitian Mandiri.docx", "storage/app/templates/Template_Laporan.docx")
prep_template("Template Proposal penelitian mandiri - Copy.docx", "storage/app/templates/Template_Proposal.docx")
print("Templates prepped successfully.")
